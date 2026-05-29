# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trPr.append(trHeight)

doc = Document()

# 页面边距
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 标题
h = doc.add_heading('4.1 周期上报数据内容', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 正文说明
for text in [
    '（1）发送ID=1，接收ID=10。',
    '（2）数据帧总计数为1，当前帧计数为0，长度为32字节。',
    '（3）周期上报数据消息ID为：0x1502000。',
]:
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()

# 表4标题
t4_title = doc.add_paragraph('表4  周期上报数据格式')
t4_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t4_title.runs[0].font.bold = True
t4_title.runs[0].font.size = Pt(11)

# 表4数据
table1_data = [
    ('序号', '名称', '字长', '数据类型', '物理量范围', '比例尺/备注'),
    ('1', '头字', '1字', '0A00H', '0~65535', ''),
    ('2', '指令编码', '1字', '无符号整型', '0~65535',
     '收到开机/关机指令后——\n0x0A07：可控整流开机；\n0x0A08：可控整流关机；\n收到开机/关机指令前——\n0x0000：空闲；\n其他：非法'),
    ('3', '命令执行状态', '1字', '无符号整型', '0~65535',
     '1校验和错；2头字错；3数据非法；4命令已经收到（正在执行）；5命令执行完毕。'),
    ('4', '电机转速', '1字', '无符号整型', '0~50000rpm', '1'),
    ('5', '母线电流', '1字', '无符号整型', '0~1000A', '电流=数字量×0.02'),
    ('6', '母线电压', '1字', '无符号整型', '0~1000V', '电压=数字量×0.02'),
    ('7', '控制器温度值', '1字', '无符号整型', '-50℃~300℃', '温度=数字量×350/65535-50'),
    ('8', '心跳', '1字', '无符号整型', '0~65535', '1'),
    ('9', '运行状态', '1字', '无符号整型', '0~65535', '1'),
    ('10', '电机温度值', '1字', '无符号整型', '-50℃~300℃', '温度=数字量×350/65535-50'),
    ('11', '电机电流幅值', '1字', '无符号整型', '0~1000A', '电流=数字量×0.02'),
    ('12', '电机功率', '1字', '无符号整型', '0~1000kw', '功率=数字量×0.02'),
    ('13', '故障码', '1字', '无符号整型', '0~65535', '1'),
    ('14', '可控整流状态字', '1字', '无符号整型', '0~65535', '1'),
    ('15', '保留', '1字', '无符号整型', '0~65535', '1'),
    ('16', '校验和', '1字', '无符号整型', '0~65535', '符合1188A校验'),
]

t1 = doc.add_table(rows=len(table1_data), cols=6)
t1.style = 'Table Grid'
t1.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Cm(1.0), Cm(2.8), Cm(1.2), Cm(2.2), Cm(2.5), Cm(5.5)]

for i, row_data in enumerate(table1_data):
    row = t1.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.width = col_widths[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 5 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(cell_text)
        run.font.size = Pt(10)
        if i == 0:
            run.font.bold = True
            set_cell_bg(cell, 'D9D9D9')

doc.add_paragraph()

# 故障码说明
p1 = doc.add_paragraph('可控整流双发控制器故障码：')
p1.runs[0].font.size = Pt(11)

p2 = doc.add_paragraph(
    'D15～D0（最低位）：均为1表示可控整流双发控制器正常；D9～D5、D2：备用。')
p2.runs[0].font.size = Pt(11)

doc.add_paragraph()

t2_title = doc.add_paragraph('可控整流双发控制器故障码位定义')
t2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2_title.runs[0].font.bold = True
t2_title.runs[0].font.size = Pt(11)

# 表2数据
table2_data = [
    ('位号', '位值为1的含义', '位值为0的含义'),
    ('D15', '控制器母线欠压故障', '控制器母线电压正常'),
    ('D14', '控制器硬件过流',     '控制器硬件不过流'),
    ('D13', '控制器驱动故障',     '控制器驱动正常'),
    ('D12', '自检故障',           '自检正常'),
    ('D11', '控制器过温故障',     '控制器温度正常'),
    ('D10', '电机过温故障',       '电机温度正常'),
    ('D9',  '备用（填0）',        '备用（填0）'),
    ('D8',  '备用（填0）',        '备用（填0）'),
    ('D7',  '备用（填0）',        '备用（填0）'),
    ('D6',  '备用（填0）',        '备用（填0）'),
    ('D5',  '备用（填0）',        '备用（填0）'),
    ('D4',  '控制器母线过压故障', '控制器母线电压正常'),
    ('D3',  '控制器软件过流',     '控制器软件不过流'),
    ('D2',  '备用（填0）',        '备用（填0）'),
    ('D1',  '控制器24V欠压故障',  '控制器24V电压正常'),
    ('D0',  '电机超速保护',       '电机速度正常'),
]

t2 = doc.add_table(rows=len(table2_data), cols=3)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

col2_widths = [Cm(1.5), Cm(6.5), Cm(6.5)]

for i, row_data in enumerate(table2_data):
    row = t2.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.width = col2_widths[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text)
        run.font.size = Pt(10)
        if i == 0:
            run.font.bold = True
            set_cell_bg(cell, 'D9D9D9')

output_path = r'c:\UserWorx\VSWorx\ServoStudio\wpfui-servoStudio\docs\周期上报数据.docx'
doc.save(output_path)
print(f'已生成：{output_path}')
